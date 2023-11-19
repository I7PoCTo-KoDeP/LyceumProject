from docx import Document
from constants import (LEARNING_HEADER, LEARNING_CRITERIA, MULTIPLIER, PER_PERFECT_STUDENT, PER_GOOD_STUDENT,
                       ACHIEVEMENT_HEADER)


class CreateWordFile:                             # Класс отвечающий за генерацию docx-файла.
    def __init__(self, database):
        self.document = Document()
        self.database = database
        self.total = 0

    def docx_file_generator(self, n, link_on_table):                    # Основная функция создания docx-файла.
        if n == 2:                                                      # Выыбор системы обучения.
            edu_sys = ['1 полугодие', '2 полугодие']
        else:
            edu_sys = ['1 четверть', '2 четверть', '3 четверть', '4 четверть']
        number_of_class = self.database.send_request('''SELECT Num_of_class FROM Other_data''')[0][0]
        year = self.database.send_request('''SELECT Date FROM Other_data''')[0][0]
        self.document.add_heading(f'Достижения {number_of_class} за 20{year} учебный год', 0)
        for i in edu_sys:
            self.document.add_heading(f'Достижения за {i}', 1)
            self.make_grades_table(i)                                       # Создание таблицы успеваемости.
            self.make_achievement_tables(i)                                 # Создание таблицы достижений.
            self.document.add_page_break()                                  # Разрыв страницы для более удобного чтения.
        self.document.add_heading(f'Резултаты за {year} учебный год', 2)    # Добавление итоговой таблицы.
        self.make_final_table(link_on_table, edu_sys)
        self.document.add_heading(f'Итого: {self.total}', 1)
        self.document.save(f'Достижения_{number_of_class}.docx')

    def make_grades_table(self, date):                  # Функция отвечающая за создание таблиц успеваемости.
        total_per_aspect = 0                            # Все данные берутся из БД и из них формируетяс таблица.
        data = self.database.send_request('''SELECT AverageScore, PerfectStudents, GoodStudents 
                                      FROM Grade WHERE Date = ?''', (date,))
        self.document.add_heading('Учебная деятельность', 2)
        table = self.document.add_table(rows=1, cols=len(LEARNING_HEADER))
        table.style = 'TableGrid'
        header_cells = table.rows[0].cells
        for i in range(len(LEARNING_HEADER)):
            header_cells[i].text = LEARNING_HEADER[i]
        if not data:
            return
        for i in range(len(LEARNING_CRITERIA)):
            new_row = table.add_row().cells
            new_row[0].text = LEARNING_CRITERIA[i]
            new_row[1].text = str(data[0][i])
            if i == 0:
                new_row[3].text = str(int(data[0][i] * MULTIPLIER))
                total_per_aspect += int(data[0][i] * MULTIPLIER)
            elif i == 1:
                if data[0][i] != '':
                    new_row[3].text = str(len(data[0][i].split(', ')) * PER_PERFECT_STUDENT)
                    total_per_aspect += len(data[0][i].split(', ')) * PER_PERFECT_STUDENT
            else:
                if data[0][i] != '':
                    new_row[3].text = str(len(data[0][i].split(', ')) * PER_GOOD_STUDENT)
                    total_per_aspect += len(data[0][i].split(', ')) * PER_GOOD_STUDENT
        self.total += total_per_aspect
        new_row = table.add_row().cells
        new_row[0].text = 'Итого'
        new_row[1].text = str(total_per_aspect)

    def make_achievement_tables(self, date):            # Функция отвечающая за создание таблиц достижений
        aspects = self.database.send_request('''SELECT Name FROM Aspects''')
        for i in aspects:                               # Все данные берутся из БД и из них формируетяс таблица.
            total_per_aspect = 0
            self.document.add_heading(i, 2)
            data = self.database.send_request('''SELECT
                                                Aspects.Name as AspectName,
                                                Type.Name as Type,
                                                Levels.Name as LevelName,
                                                Places.Name as Place,
                                                MainTable.Name,
                                                MainTable.Participants,
                                                Type.Points
                                            FROM
                                                MainTable
                                            LEFT JOIN Aspects ON MainTable.AspectId = Aspects.Id
                                            LEFT JOIN Type ON MainTable.TypeId = Type.Id
                                            LEFT JOIN Levels ON Type.LevelId = Levels.Id
                                            LEFT JOIN Places ON Type.PlaceId = Places.Id
                                            WHERE AspectName = ? AND Date = ?
                                            ORDER BY AspectName;''', (*i, date))
            table = self.document.add_table(rows=1, cols=len(ACHIEVEMENT_HEADER) - 1)
            table.style = 'TableGrid'
            header_cells = table.rows[0].cells
            for j in range(len(ACHIEVEMENT_HEADER) - 1):
                header_cells[j].text = ACHIEVEMENT_HEADER[j + 1]
            for row in data:
                new_row = table.add_row().cells
                students = 0
                for n, elem in enumerate(row):
                    if n != 0:
                        new_row[n - 1].text = str(elem)
                    if n == 5:
                        students = len(elem.split(', '))
                    elif n == 6:
                        total_per_aspect += students * int(elem)
            self.total += total_per_aspect
            new_row = table.add_row().cells
            new_row[0].text = 'Итого'
            new_row[1].text = str(total_per_aspect)

    def make_final_table(self, link_on_table, dates):       # Функция отвечающая за создание итоговой таблицы.
        table = self.document.add_table(rows=1, cols=link_on_table.columnCount())
        table.style = 'TableGrid'                           # link_on_table - ссылка на таблицу в основном окне
        header_cells = table.rows[0].cells                  # оттуда берутся все данные, вместо того чтобы их считать
        for i in range(len(dates) + 1):                     # снова.
            if i == 0:
                header_cells[i].text = ''
            else:
                header_cells[i].text = dates[i - 1]
        for i in range(link_on_table.rowCount()):
            new_row = table.add_row().cells
            for j in range(link_on_table.columnCount()):
                if link_on_table.item(i, j) is not None:
                    new_row[j].text = link_on_table.item(i, j).text()
