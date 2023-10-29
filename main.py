import sys
import sqlite3
from docx import Document
from PyQt5 import uic
from PyQt5.QtWidgets import (QApplication, QMainWindow, QTableWidgetItem, QAction, QWidget, QMenu, QDialog,
                             QMessageBox, QListWidgetItem)

ACHIEVEMENT = 1
LEARNING = 0
MULTIPLIER = 100
PER_GOOD_STUDENT = 30
PER_PERFECT_STUDENT = 50
DATABASE_PATH = 'data/auto-sys-database.sqlite'
ACHIEVEMENT_HEADER = ['Вид', 'Тип', 'Уровень', 'Место', 'Описание', 'Участники', 'Очки']
LEARNING_HEADER = ['Критерии', 'Кол-во', 'Ср.балл, кол-во', 'Очки']
LEARNING_CRITERIA = ['Ср. Балл по успеваемости', 'Отличники', 'Ударники']
MAIN_CRITERIA = ['Учёба', 'Внеурочная деятельность', 'Спорт', 'Активная жизненная позиция']


class AchievementControl(QMainWindow):
    def __init__(self):
        super().__init__()
        self.roster = {}
        self.database = Database(DATABASE_PATH)
        self.docx = WorkingWithWordFiles(self.database)
        # Load UI.
        uic.loadUi('UIs/cadet.ui', self)
        self.initUI()

    def initUI(self):
        menubar = self.menuBar()
        # Actions.
        makefile_action = QAction('&Сгенерировать файл', self)
        makefile_action.triggered.connect(self.create_docx_file)
        save_action = QAction('&Сохранить', self)
        about_action = QAction('&О программе', self)
        about_action.triggered.connect(self.open_info)
        # Menus.
        file_menu = menubar.addMenu('&Файл')
        workspace_menu = menubar.addMenu('&Открыть таблицу')
        info_menu = menubar.addMenu('&Справка')
        # Submenus.
        self.learning_menu = QMenu('&Учёба', self)
        self.achievement_menu = QMenu('&Достижения', self)
        # Fill submenus.
        self.login()
        # Another methods.
        info_menu.addAction(about_action)
        file_menu.addAction(save_action)
        file_menu.addAction(makefile_action)
        workspace_menu.addMenu(self.learning_menu)
        workspace_menu.addMenu(self.achievement_menu)
        # Initialization methods.
        self.load_table()
        self.show_best_students()

    def quarters(self):
        self.n = 4
        for i in range(1, 5):
            learning_action = QAction(f'&Открыть {i} четверть', self)
            achievement_action = QAction(f'&Открыть {i} четверть', self)
            learning_action.setData(LEARNING)
            achievement_action.setData(ACHIEVEMENT)
            learning_action.triggered.connect(self.open_workspace)
            achievement_action.triggered.connect(self.open_workspace)
            self.learning_menu.addAction(learning_action)
            self.achievement_menu.addAction(achievement_action)

    def half_year(self):
        self.n = 2
        for i in range(1, 3):
            learning_action = QAction(f'&Открыть {i} полугодие', self)
            achievement_action = QAction(f'&Открыть {i} полугодие', self)
            learning_action.setData(LEARNING)
            achievement_action.setData(ACHIEVEMENT)
            learning_action.triggered.connect(self.open_workspace)
            achievement_action.triggered.connect(self.open_workspace)
            self.learning_menu.addAction(learning_action)
            self.achievement_menu.addAction(achievement_action)

    def login(self):
        with open('data/settings_file.txt', mode='r', encoding='utf-8') as f:
            if int(f.read(2)) >= 10:
                self.half_year()
            else:
                self.quarters()

    def load_table(self):
        total = 0
        total_for_piece = [0] * self.n
        self.tableWidget.setColumnCount(self.n + 1)
        self.tableWidget.setRowCount(4)
        if self.n == 2:
            self.tableWidget.setHorizontalHeaderLabels(['', '1 полугодие', '2 полугодие'])
        else:
            self.tableWidget.setHorizontalHeaderLabels(['', '1 четверть', '2 четверть', '3 четверть', '4 четверть'])
        for i in range(len(MAIN_CRITERIA)):
            for j in range(self.n + 1):
                points = 0
                if j == 0:
                    self.tableWidget.setItem(i, j, QTableWidgetItem(MAIN_CRITERIA[i]))
                if i == 0 and j != 0:
                    points = self.calculate_learning(j)
                    self.tableWidget.setItem(i, j, QTableWidgetItem(str(points)))
                if i == 1 and j != 0:
                    points = self.calculate_achievements(j, MAIN_CRITERIA[i])
                    self.tableWidget.setItem(i, j, QTableWidgetItem(str(points)))
                if i == 2 and j != 0:
                    points = self.calculate_achievements(j, MAIN_CRITERIA[i])
                    self.tableWidget.setItem(i, j, QTableWidgetItem(str(points)))
                if i == 3 and j != 0:
                    points = self.calculate_achievements(j, MAIN_CRITERIA[i])
                    self.tableWidget.setItem(i, j, QTableWidgetItem(str(points)))
                total_for_piece[j - 1] += points
                total += points
        self.tableWidget.setRowCount(self.tableWidget.rowCount() + 1)
        self.tableWidget.setItem(self.tableWidget.rowCount() - 1, 0, QTableWidgetItem('Итого'))
        for i, val in enumerate(total_for_piece):
            self.tableWidget.setItem(self.tableWidget.rowCount() - 1, i + 1, QTableWidgetItem(str(int(val))))
        self.tableWidget.setRowCount(self.tableWidget.rowCount() + 1)
        self.tableWidget.setItem(self.tableWidget.rowCount() - 1, 0, QTableWidgetItem('За учебный год'))
        self.tableWidget.setItem(self.tableWidget.rowCount() - 1, 1, QTableWidgetItem(str(int(total))))
        self.tableWidget.resizeColumnsToContents()

    def open_workspace(self):
        self.workspace = WorkspaceWindow(self.sender().text()[self.sender().text().find(' ') + 1:],
                                         self.sender().data(), self.database)
        self.workspace.show()

    def open_info(self):
        self.info_menu = Info()
        self.info_menu.show()

    def calculate_learning(self, date):
        result = 0
        learning = self.database.get_data('''SELECT * FROM Grade''')
        for avg_score, perf_stud, good_stud, dt in learning:
            if int(dt[0]) == date:
                n, k = len(perf_stud.split(', ')), len(good_stud.split(', '))
                if perf_stud == '':
                    n = 0
                if good_stud == '':
                    k = 0
                result = int(avg_score * MULTIPLIER + n * PER_PERFECT_STUDENT + k * PER_GOOD_STUDENT)
        if result is None:
            return 0
        return result

    def calculate_achievements(self, date, aspect):
        result = 0
        achievements = self.database.get_data('''SELECT
                                              Aspects.Name as AspectName,
                                              MainTable.Participants,
                                              Type.Points,
                                              MainTable.Date
                                              FROM
                                              MainTable
                                              LEFT JOIN Aspects ON MainTable.AspectId = Aspects.Id
                                              LEFT JOIN Type ON MainTable.TypeId = Type.Id
                                              ORDER BY AspectName;''')
        for asp, participant, points, dt in achievements:
            if asp == aspect and int(dt[0]) == date:
                result += len(participant.split(', ')) * points
        if result is None:
            return 0
        return result

    def show_best_students(self):
        learning = self.database.get_data('''SELECT * FROM Grade''')
        for i in learning:
            for j in i[1].split(', '):
                if j not in self.roster:
                    self.roster[j] = PER_PERFECT_STUDENT
                else:
                    self.roster[j] += PER_PERFECT_STUDENT
            for j in i[2].split(', '):
                if j not in self.roster:
                    self.roster[j] = PER_GOOD_STUDENT
                else:
                    self.roster[j] += PER_GOOD_STUDENT
        achievements = self.database.get_data('''SELECT MainTable.Participants, Type.Points FROM MainTable
                                              LEFT JOIN Type ON MainTable.TypeId = Type.Id''')
        for i in achievements:
            for j in i[0].split(', '):
                if j in self.roster:
                    self.roster[j] += i[1]
                else:
                    self.roster[j] = i[1]
        sorted_roster = sorted(self.roster.items(), key=lambda x: x[1], reverse=True)
        for i, j in enumerate(sorted_roster):
            if j[0] != '':
                self.listWidget.addItem(QListWidgetItem(j[0].ljust(15, ' ') + str(j[1])))

    def create_docx_file(self):
        self.docx.docx_file_generator(self.n, self.tableWidget)

    def closeEvent(self, event):
        self.database.close_connection()


class WorkspaceWindow(QMainWindow):
    def __init__(self, date, aspect, database):
        self.date = date
        self.aspect = aspect
        self.database = database
        self.is_changed = False

        super().__init__()
        if aspect == LEARNING:
            uic.loadUi('UIs/learning_workspace.ui', self)
        else:
            uic.loadUi('UIs/workspace.ui', self)
            self.initUI()
        self.load_table(self.aspect)

    def initUI(self):
        self.setWindowTitle(f'{self.date} - Рабочая область')
        menubar = self.menuBar()
        self.statusbar = self.statusBar()
        # Actions.
        find_action = QAction('&Найти', self)
        delete_action = QAction('Удалить выбранную строку', self)
        delete_action.triggered.connect(self.delete_row)
        # Menus.
        redo_menu = menubar.addMenu('&Правка')

        redo_menu.addAction(find_action)
        redo_menu.addAction(delete_action)
        self.add_button.clicked.connect(self.put_in_table)
        self.load_combo_boxes()

    def load_table(self, aspect):                                           # Вывод таблицы.
        num_of_students = 0                                                 # Также можно используется, как обновление
        total = 0                                                           # данных в таблице.
        points = 0
        res = self.database.get_table(self.date, self.aspect)
        self.data_output.setRowCount(0)
        if aspect == ACHIEVEMENT:
            if not res:
                return
            self.data_output.setColumnCount(len(res[0]) - 1)
            self.data_output.setHorizontalHeaderLabels(ACHIEVEMENT_HEADER)
            for i, row in enumerate(res):
                self.data_output.setRowCount(self.data_output.rowCount() + 1)
                for j, elem in enumerate(row):
                    if j == 7:
                        self.data_output.item(i, 4).setData(20, elem)
                    elif j == 5:
                        num_of_students = len(elem.split(', '))
                        self.data_output.setItem(i, j, QTableWidgetItem(str(elem)))
                    elif j == 6:
                        self.data_output.setItem(i, j, QTableWidgetItem(str(elem * num_of_students)))
                        total += elem * num_of_students
                    else:
                        self.data_output.setItem(i, j, QTableWidgetItem(str(elem)))
        else:
            res = res[0]
            self.data_output.setColumnCount(len(LEARNING_HEADER))
            self.data_output.setHorizontalHeaderLabels(LEARNING_HEADER)
            for i in range(len(LEARNING_CRITERIA)):
                self.data_output.setRowCount(self.data_output.rowCount() + 1)
                for j in range(len(LEARNING_HEADER)):
                    if j == 0:
                        self.data_output.setItem(i, j, QTableWidgetItem(str(LEARNING_CRITERIA[i])))
                    elif j == 1 and i > 0:
                        self.data_output.setItem(i, j, QTableWidgetItem(str(res[i])))
                        if res[i] != '':
                            num_of_students = len(res[i].split(', '))
                        else:
                            num_of_students = 0
                        if i == 1:
                            points = num_of_students * PER_PERFECT_STUDENT
                        elif i == 2:
                            points = num_of_students * PER_GOOD_STUDENT
                        total += points
                    elif j == 2 and i == 0:
                        self.data_output.setItem(i, j, QTableWidgetItem(str(res[i])))
                        points = float(res[i]) * MULTIPLIER
                        total += points
                    elif j == 2 and i > 0:
                        self.data_output.setItem(i, j, QTableWidgetItem(str(num_of_students)))
                    elif j == 3 and i > 0:
                        self.data_output.setItem(i, j, QTableWidgetItem(str(points)))
                    elif j == 3 and i == 0:
                        self.data_output.setItem(i, j, QTableWidgetItem(str(int(points))))
        self.data_output.setRowCount(self.data_output.rowCount() + 1)
        self.data_output.setItem(self.data_output.rowCount() - 1, 0, QTableWidgetItem('Итого'))
        self.data_output.setItem(self.data_output.rowCount() - 1, 1, QTableWidgetItem(str(int(total))))
        self.data_output.resizeColumnsToContents()
        self.data_output.cellChanged.connect(self.cell_changed)

    def load_combo_boxes(self):                                         # Загрузка интерфейсов comboBoxes.
        type_data = set(self.database.get_data('''SELECT Name FROM Type'''))
        aspect_data = self.database.get_data('''SELECT Name FROM Aspects''')
        level_data = self.database.get_data('''SELECT Name FROM LevelId''')
        place_data = self.database.get_data('''SELECT Name FROM Places''')
        for i in aspect_data:
            self.aspect_box.addItem(*i)
        for i in type_data:
            self.type_box.addItem(*i)
        for i in place_data:
            self.place_box.addItem(*i)
        for i in level_data:
            self.level_box.addItem(*i)
        self.aspect_box.textActivated.connect(self.change_types)        # При изменении значения в одном ComboBox
        self.type_box.textActivated.connect(self.change_place)          # меняются значения и в других.
        self.type_box.textActivated.connect(self.change_level)

    def change_types(self, text):
        self.type_box.clear()
        aspect_id = self.database.get_data('''SELECT Id FROM Aspects WHERE Name = ?''', (text,))
        type_data = set(self.database.get_data('''SELECT Name FROM Type WHERE AspectId = ?''', *aspect_id))
        for i in type_data:
            self.type_box.addItem(*i)

    def change_place(self, text):
        self.place_box.clear()
        checklist = []
        place_ids = self.database.get_data('''SELECT PlaceId FROM Type WHERE Name = ?''', (text,))
        for i in place_ids:
            place = self.database.get_data('''SELECT Name FROM Places WHERE Id = ?''', i)[0]
            if place[0] not in checklist:
                checklist.append(*place)
                self.place_box.addItem(*place)

    def change_level(self, text):
        self.level_box.clear()
        checklist = []
        level_ids = self.database.get_data('''SELECT LevelId FROM Type WHERE Name = ?''', (text,))
        for i in level_ids:
            level = self.database.get_data('''SELECT Name FROM LevelId WHERE Id = ?''', i)
            if level:
                level = level[0][0]
                if level not in checklist:
                    checklist.append(level)
                    self.level_box.addItem(level)
            else:
                return

    def put_in_table(self):                                             # Добавление ряда в базу данных.
        self.is_changed = True                                          # Работает только с достижениями
        if self.aspect == ACHIEVEMENT:
            request = [self.aspect_box.currentText(), self.type_box.currentText(), self.place_box.currentText(),
                       self.level_box.currentText(), self.request_edit.toPlainText()]
            self.database.add_to_database(self.aspect, self.date, *request)
            self.load_table(self.aspect)

    def change_data(self):                                              # Изменение данных.
        if self.aspect == ACHIEVEMENT:                                  # Берётся информация из каждого ряда
            for i in range(self.data_output.rowCount() - 1):            # и перезаписывается
                asp = self.data_output.item(i, 0).text()
                type = self.data_output.item(i, 1).text()
                level = self.data_output.item(i, 2).text()
                place = self.data_output.item(i, 3).text()
                id = self.data_output.item(i, 4).data(20)
                name = self.data_output.item(i, 4).text()
                participants = self.data_output.item(i, 5).text()
                self.database.edit_data(self.aspect, self.date, id, asp, type, level, place, name, participants)
        else:
            avg_points = self.data_output.item(0, 2).text()
            perf_stud = self.data_output.item(1, 1).text()
            good_stud = self.data_output.item(2, 1).text()
            self.database.edit_data(self.aspect, self.date, avg_points, perf_stud, good_stud)

    def delete_row(self):
        self.is_changed = True
        id = self.data_output.item(self.data_output.currentRow(), 4).data(20)
        self.database.delete_data(id)
        self.load_table(self.aspect)

    def cell_changed(self, item):
        self.is_changed = True

    def open_find_window(self):
        self.find_window = FindWindow()
        self.find_window.show()

    def closeEvent(self, event):
        if self.is_changed:
            valid = QMessageBox.question(self, '', 'Сохранить изменения?', QMessageBox.Yes, QMessageBox.No)
            if valid == QMessageBox.Yes:
                self.change_data()
                self.database.confirm_changes()


class RegistrationWindow(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('UIs/register.ui', self)
        self.save_button.clicked.connect(self.save)

    def save(self):
        with open('data/settings_file.txt', mode='w', encoding='utf-8') as f:
            f.write(self.class_edit.text() + ' ' + self.platoon_edit.text())
            self.accept()
            self.close()


class FindWindow(QWidget):
    def __init__(self):
        super().__init__()
        uic.loadUi('UIs/Info.ui', self)


class Info(QWidget):
    def __init__(self):
        super().__init__()
        uic.loadUi('UIs/Info.ui', self)
        self.close_button.clicked.connect(self.close)


class Database:
    def __init__(self, path):
        self.connection = sqlite3.connect(path)
        self.cursor = self.connection.cursor()

    def get_table(self, date, aspect):
        if aspect == ACHIEVEMENT:
            return self.cursor.execute('''SELECT
                                                    Aspects.Name as AspectName,
                                                    Type.Name as Type,
                                                    LevelId.Name as LevelName,
                                                    Places.Name as Place,
                                                    MainTable.Name,
                                                    MainTable.Participants,
                                                    Type.Points,
                                                    MainTable.Id
                                                FROM
                                                    MainTable
                                                LEFT JOIN Aspects ON MainTable.AspectId = Aspects.Id
                                                LEFT JOIN Type ON MainTable.TypeId = Type.Id
                                                LEFT JOIN LevelId ON Type.LevelId = LevelId.Id
                                                LEFT JOIN Places ON Type.PlaceId = Places.Id
                                                WHERE Date = ?
                                                ORDER BY AspectName;''', (date,)).fetchall()
        else:
            return self.cursor.execute('''SELECT AverageScore, PerfectStudents, GoodStudents FROM Grade
                                            WHERE Date = ?''', (date,)).fetchall()

    def get_data(self, request, *args):
        if args:
            return self.cursor.execute(request, *args).fetchall()
        return self.cursor.execute(request).fetchall()

    def add_to_database(self, aspect, date, aspect_name, type_name, place_name, level_name, other):
        if aspect == ACHIEVEMENT:
            other = other.split('; ')
            place_id = self.cursor.execute('''SELECT Id FROM Places WHERE Name = ?''', (place_name,)).fetchone()
            level_id = self.cursor.execute('''SELECT Id FROM LevelId WHERE Name = ?''', (level_name,)).fetchone()
            aspect_id = self.cursor.execute('''SELECT Id FROM Aspects WHERE Name = ?''', (aspect_name,)).fetchone()
            type_id = self.cursor.execute('''SELECT Id FROM Type WHERE Name = ? AND PlaceId = ? AND AspectId = ? AND 
                                          LevelId = ?''', (type_name, *place_id, *aspect_id, *level_id)).fetchone()
            self.cursor.execute('''INSERT INTO MainTable VALUES(?, ?, ?, ?, null, ?)''',
                                (*aspect_id, other[0], other[1], *type_id, date)).fetchone()

    def edit_data(self, aspect, date, *new_row):
        if aspect == ACHIEVEMENT:
            id, aspect_name, type_name, level, place_name, name, stud = new_row
            level_id = self.cursor.execute('''SELECT Id FROM LevelId WHERE Name = ?''', (level,)).fetchone()
            place_id = self.cursor.execute('''SELECT Id FROM Places WHERE Name = ?''', (place_name,)).fetchone()
            aspect_id = self.cursor.execute('''SELECT Id FROM Aspects WHERE Name = ?''', (aspect_name,)).fetchone()
            type_id = self.cursor.execute('''SELECT Id FROM Type WHERE Name = ? AND PlaceId = ? AND AspectId = ? AND 
                                          LevelId = ?''', (type_name, *place_id, *aspect_id, *level_id)).fetchone()
            self.cursor.execute('''UPDATE MainTable
                                SET AspectId = ?, Name = ?, Participants = ?, TypeId = ? WHERE Id = ?''',
                                (*aspect_id, name, stud, *type_id, id)).fetchall()
        else:
            avg, perf, good = new_row
            self.cursor.execute('''UPDATE Grade
                                SET AverageScore = ?, PerfectStudents = ?, GoodStudents = ?
                                WHERE Date = ?''', (avg, perf, good, date)).fetchall()

    def delete_data(self, row_id):
        self.cursor.execute('''DELETE FROM MainTable WHERE id = ?''', (row_id,)).fetchall()

    def confirm_changes(self):
        self.connection.commit()

    def close_connection(self):
        self.connection.close()


class WorkingWithWordFiles:                             # Класс отвечающий за генерацию docx-файла.
    def __init__(self, database):
        self.document = Document()
        self.database = database
        self.total = 0

    def docx_file_generator(self, n, link_on_table):
        if n == 2:
            edu_sys = ['1 полугодие', '2 полугодие']
        else:
            edu_sys = ['1 четверть', '2 четверть', '3 четверть', '4 четверть']
        with open('data/settings_file.txt', mode='r', encoding='utf-8') as f:
            number_of_class = f.read().split()[1]
        self.document.add_heading(f'Достижения {number_of_class}', 0)
        for i in edu_sys:
            self.document.add_heading(f'Достижения за {i}', 1)
            self.make_grades_table(i)
            self.make_achievement_tables(i)
            self.document.add_page_break()
        self.make_final_table(link_on_table, edu_sys)
        self.document.add_heading(f'Итого: {self.total}', 1)
        self.document.save(f'Достижения_{number_of_class}.docx')

    def make_grades_table(self, date):                  # Создание таблиц успеваемости.
        total_per_aspect = 0
        data = self.database.get_data('''SELECT AverageScore, PerfectStudents, GoodStudents 
                                      FROM Grade WHERE Date = ?''', (date,))
        self.document.add_heading('Учебная деятельность', 2)
        table = self.document.add_table(rows=1, cols=len(LEARNING_HEADER))
        table.style = 'TableGrid'
        header_cells = table.rows[0].cells
        for i in range(len(LEARNING_HEADER)):
            header_cells[i].text = LEARNING_HEADER[i]
        if not data:
            return
        for i in range(len(LEARNING_CRITERIA)):
            new_row = table.add_row().cells
            new_row[0].text = LEARNING_CRITERIA[i]
            new_row[1].text = str(data[0][i])
            if i == 0:
                new_row[3].text = str(int(data[0][i] * MULTIPLIER))
                total_per_aspect += int(data[0][i] * MULTIPLIER)
            elif i == 1:
                if data[0][i] != '':
                    new_row[3].text = str(len(data[0][i].split(', ')) * PER_PERFECT_STUDENT)
                    total_per_aspect += len(data[0][i].split(', ')) * PER_PERFECT_STUDENT
            else:
                if data[0][i] != '':
                    new_row[3].text = str(len(data[0][i].split(', ')) * PER_GOOD_STUDENT)
                    total_per_aspect += len(data[0][i].split(', ')) * PER_GOOD_STUDENT
        self.total += total_per_aspect
        new_row = table.add_row().cells
        new_row[0].text = 'Итого'
        new_row[1].text = str(total_per_aspect)

    def make_achievement_tables(self, date):            # Создание таблиц достижений
        aspects = self.database.get_data('''SELECT Name FROM Aspects''')
        for i in aspects:
            total_per_aspect = 0
            self.document.add_heading(i, 2)
            data = self.database.get_data('''SELECT
                                                Aspects.Name as AspectName,
                                                Type.Name as Type,
                                                LevelId.Name as LevelName,
                                                Places.Name as Place,
                                                MainTable.Name,
                                                MainTable.Participants,
                                                Type.Points
                                            FROM
                                                MainTable
                                            LEFT JOIN Aspects ON MainTable.AspectId = Aspects.Id
                                            LEFT JOIN Type ON MainTable.TypeId = Type.Id
                                            LEFT JOIN LevelId ON Type.LevelId = LevelId.Id
                                            LEFT JOIN Places ON Type.PlaceId = Places.Id
                                            WHERE AspectName = ? AND Date = ?
                                            ORDER BY AspectName;''', (*i, date))
            table = self.document.add_table(rows=1, cols=len(ACHIEVEMENT_HEADER) - 1)
            table.style = 'TableGrid'
            header_cells = table.rows[0].cells
            for j in range(len(ACHIEVEMENT_HEADER) - 1):
                header_cells[j].text = ACHIEVEMENT_HEADER[j + 1]
            for row in data:
                new_row = table.add_row().cells
                for n, elem in enumerate(row):
                    if n != 0:
                        new_row[n - 1].text = str(elem)
                    if n == 5:
                        students = len(elem.split(', '))
                    elif n == 6:
                        total_per_aspect += students * int(elem)
            self.total += total_per_aspect
            new_row = table.add_row().cells
            new_row[0].text = 'Итого'
            new_row[1].text = str(total_per_aspect)

    def make_final_table(self, link_on_table, dates):
        table = self.document.add_table(rows=1, cols=link_on_table.columnCount())
        table.style = 'TableGrid'
        header_cells = table.rows[0].cells
        for i in range(len(dates) + 1):
            if i == 0:
                header_cells[i].text = ''
            else:
                header_cells[i].text = dates[i - 1]
        for i in range(link_on_table.rowCount()):
            new_row = table.add_row().cells
            for j in range(link_on_table.columnCount()):
                if link_on_table.item(i, j) is not None:
                    new_row[j].text = link_on_table.item(i, j).text()


def registered():
    try:
        with open('data/settings_file.txt', mode='r', encoding='utf-8') as f:
            if f.read(1) == '':
                return False
    except FileNotFoundError:
        return False
    return True


if __name__ == '__main__':
    app = QApplication(sys.argv)
    if not registered():
        login = RegistrationWindow()
        if login.exec_() == QDialog.Accepted:
            ex = AchievementControl()
            ex.show()
            sys.exit(app.exec_())
    else:
        ex = AchievementControl()
        ex.show()
        sys.exit(app.exec_())
